import tempfile
import unittest
import asyncio
import base64
from io import BytesIO
from pathlib import Path
from unittest.mock import Mock, patch

from osint.dorks import DorkGenerator
from osint.evidence_capture import (
    SERPEvidenceCapturePipeline,
    classify_page_access,
    group_evidence_positions,
    sanitize_path_component,
)
from osint.models import CollectorResult, DorkQuery, EvidenceScreenshotRecord, Observation, PageCaptureRecord, SearchResult
from osint.normalizer import DomainNormalizer, TargetNormalizationError
from osint.risk import RiskScorer
from osint.storage import OSINTRepository
from docx import Document

from osint.documents import assess_docx, assess_pdf
from core.playwright_session import close_browser_session, launch_browser_session
from osint.search import AggregatingSearchProvider, BingRSSSearchProvider, DuckDuckGoSearchProvider, GoogleSearchProvider, SearchProvider
from osint.collectors.brave_search import KeylessSearchCollector
from osint.collectors.base import CollectorContext
from osint.collectors.free_discovery import CertificateTransparencyCollector, WaybackCDXCollector
from osint.collectors.x_authenticated import post_matches_target, target_variants
from osint.collectors.trustpilot import trustpilot_profile_matches, trustpilot_target_variants
from osint.collectors.authenticated_social import (
    AUTHENTICATED_SOCIAL_COLLECTORS,
    InstagramAuthenticatedCollector,
    configured_social_session_path,
)
from osint.cancellation import InvestigationCancelled
from osint.docx_report import OSINTDocxReportBuilder
from osint.text_cleanup import clean_evidence_text, evidence_scope
from ui.osint_workspace import BRAND_SCOPED_COLLECTORS, collectors_for_domain


class DomainNormalizerTests(unittest.TestCase):
    def test_normalizes_url_to_ascii_domain(self):
        target = DomainNormalizer.normalize("https://WWW.Example.com/path?q=1")
        self.assertEqual(target.domain, "www.example.com")
        self.assertEqual(target.url, "https://www.example.com")

    def test_rejects_ip_address(self):
        with self.assertRaises(TargetNormalizationError):
            DomainNormalizer.normalize("127.0.0.1")


class XAuthenticatedCollectorTests(unittest.TestCase):
    def test_matches_brand_or_domain_in_post_text(self):
        normalized = DomainNormalizer.normalize("parimatch.com")
        target = type(normalized)(
            normalized.raw_input, normalized.domain, normalized.url, brand="Parimatch"
        )
        self.assertIn("parimatch", target_variants(target))
        self.assertEqual(post_matches_target("Complaint about @Parimatch support", target), "parimatch")
        self.assertEqual(post_matches_target("Unrelated gaming discussion", target), "")

    def test_brand_inside_tracking_url_is_not_a_social_finding(self):
        normalized = DomainNormalizer.normalize("1xbet.com")
        target = type(normalized)(normalized.raw_input, normalized.domain, normalized.url, brand="1xbet")
        self.assertEqual(
            post_matches_target("Generic bonus https://tracker.example/all-casinos-1xbet-offers", target),
            "",
        )

    def test_persists_authenticated_x_screenshot_paths(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = OSINTRepository(Path(temp_dir) / "osint.db")
            target = DomainNormalizer.normalize("parimatch.com")
            repository.create_investigation("OSINT_X", target)
            screenshot = Path(temp_dir) / "x-scroll.png"
            screenshot.write_bytes(b"png")
            repository.save_collector_result(
                "OSINT_X",
                CollectorResult("x_authenticated_playwright", "COMPLETED", [Observation(
                    "x_authenticated_playwright", "Automated social and review findings",
                    "AUTOMATED_SOCIAL_FINDING", "Parimatch complaint",
                    "https://x.com/example/status/123", confidence=0.9,
                    metadata={
                        "platform": "X/Twitter", "title": "Parimatch complaint",
                        "post_text": "Complaint about Parimatch", "matched_target_variant": "parimatch",
                        "normalized_url": "https://x.com/example/status/123",
                        "collector_method": "x_authenticated_playwright",
                        "search_engine": "x_authenticated_browser", "query_id": "X-DIRECT",
                        "rank": 1, "status": "authenticated_page_captured",
                        "screenshot_paths": [str(screenshot)],
                    },
                )]),
            )
            finding = repository.get_social_findings("OSINT_X")[0]
            self.assertEqual(finding["screenshot_paths"], [str(screenshot)])


class TrustpilotCollectorTests(unittest.TestCase):
    def test_matches_target_company_profile(self):
        normalized = DomainNormalizer.normalize("parimatch.com")
        target = type(normalized)(
            normalized.raw_input, normalized.domain, normalized.url, brand="Parimatch"
        )
        self.assertIn("parimatch.com", trustpilot_target_variants(target))
        self.assertEqual(
            trustpilot_profile_matches(
                "/review/parimatch.com", "Parimatch reviews", target
            ),
            "parimatch.com",
        )

    def test_rejects_unrelated_company_profile(self):
        target = DomainNormalizer.normalize("parimatch.com")
        self.assertEqual(
            trustpilot_profile_matches(
                "/review/unrelated.example", "Unrelated company reviews", target
            ),
            "",
        )


class AuthenticatedSocialCollectorTests(unittest.TestCase):
    def test_registers_all_saved_session_platform_collectors(self):
        self.assertEqual(
            set(AUTHENTICATED_SOCIAL_COLLECTORS),
            {
                "Instagram authenticated search",
                "Facebook authenticated search",
                "Telegram authenticated search",
                "YouTube authenticated search",
                "Quora authenticated search",
            },
        )

    def test_collector_uses_platform_session_path(self):
        collector = InstagramAuthenticatedCollector(Path("/tmp/instagram-session.json"))
        self.assertEqual(collector.platform.label, "Instagram")
        self.assertEqual(collector.name, "instagram_authenticated_playwright")
        self.assertEqual(collector.session_path, Path("/tmp/instagram-session.json"))
        self.assertEqual(
            configured_social_session_path("instagram").name,
            "instagram.json",
        )

    def test_brand_scoped_collectors_run_only_for_primary_domain(self):
        selected = [
            "DNS",
            "Keyless Web Search (no API key)",
            "X/Twitter authenticated search",
            "Instagram authenticated search",
            "Trustpilot public reviews",
        ]
        self.assertEqual(collectors_for_domain(selected, 0), selected)
        secondary = collectors_for_domain(selected, 1)
        self.assertEqual(secondary, ["DNS", "Keyless Web Search (no API key)"])
        self.assertFalse(BRAND_SCOPED_COLLECTORS.intersection(secondary))


class DorkGeneratorTests(unittest.TestCase):
    def test_generates_social_and_document_queries(self):
        target = DomainNormalizer.normalize("example.com")
        queries = DorkGenerator().generate(target)
        query_by_id = {item.query_id: item for item in queries}
        self.assertEqual(len(queries), 62)
        self.assertEqual(queries[0].query_id, "B001")
        self.assertEqual(queries[0].query, 'filetype:pdf "Example"')
        self.assertIn('"example.com" site:x.com', query_by_id["Q019"].query)
        self.assertIn("site:reddit.com", query_by_id["Q022"].query)
        self.assertEqual(query_by_id["Q033"].priority, "critical")
        self.assertIn("site:example.com filetype:pdf", query_by_id["Q032"].query)
        self.assertIn("deposit not credited", query_by_id["Q029"].evidence_keywords)
        self.assertIn("court order", query_by_id["Q041"].evidence_keywords)
        self.assertTrue(all(query.evidence_keywords for query in queries))


class RiskScorerTests(unittest.TestCase):
    def test_deduplicates_and_caps_indicators(self):
        observation = Observation(
            "test", "Applications.Direct downloads", "APPLICATION_URL", "https://example.com/app.apk",
            "https://example.com", confidence=0.9, risk_points=10,
        )
        assessment = RiskScorer.assess([observation, observation])
        self.assertEqual(assessment.score, 10)
        self.assertEqual(len(assessment.indicators), 1)


class RepositoryTests(unittest.TestCase):
    def test_persists_queries_and_deduplicated_observations(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = OSINTRepository(Path(temp_dir) / "osint.db")
            target = DomainNormalizer.normalize("example.com")
            repository.create_investigation("OSINT_TEST", target)
            repository.save_queries("OSINT_TEST", DorkGenerator().generate(target))
            observation = Observation("test", "Identity", "DOMAIN", "example.com", "https://example.com")
            repository.save_collector_result("OSINT_TEST", CollectorResult("test", "COMPLETED", [observation, observation]))

            self.assertGreater(len(repository.get_queries("OSINT_TEST")), 10)
            self.assertEqual(len(repository.get_observations("OSINT_TEST")), 1)

    def test_maps_duplicate_source_to_multiple_queries(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = OSINTRepository(Path(temp_dir) / "osint.db")
            target = DomainNormalizer.normalize("example.com")
            repository.create_investigation("OSINT_MAP", target)
            observations = []
            for query_id, rank in (("Q001", 2), ("Q029", 1)):
                observations.append(
                    Observation(
                        "brave_search", "Search.money_flow", "SEARCH_RESULT", "Example evidence",
                        "https://reports.example.org/item?utm_source=test", confidence=0.7,
                        metadata={
                            "query_id": query_id,
                            "query_text": "query",
                            "search_engine": "brave",
                            "rank": rank,
                            "title": "Example evidence",
                            "snippet": "deposit evidence",
                            "normalized_url": "https://reports.example.org/item",
                            "source_type": "unknown",
                        },
                    )
                )
            repository.save_collector_result("OSINT_MAP", CollectorResult("brave_search", "COMPLETED", observations))
            sources = repository.get_sources("OSINT_MAP")
            self.assertEqual(len(sources), 1)
            self.assertEqual(sources[0]["discovered_by_queries"], 2)
            self.assertEqual(sources[0]["best_rank"], 1)

    def test_prioritizes_document_query_sources_for_download(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = OSINTRepository(Path(temp_dir) / "osint.db")
            target = DomainNormalizer.normalize("example.com")
            repository.create_investigation("OSINT_DOCS", target)
            repository.save_queries("OSINT_DOCS", DorkGenerator().generate(target))
            observations = [
                Observation("search", "Search", "SEARCH_RESULT", "Web page", "https://example.net/page", metadata={"query_id": "B002", "query_text": "reviews", "search_engine": "test", "rank": 1, "title": "Example reviews", "snippet": "Example review", "normalized_url": "https://example.net/page", "source_type": "unknown"}),
                Observation("search", "Search", "SEARCH_RESULT", "PDF", "https://papers.test/example.pdf", metadata={"query_id": "B001", "query_text": "filetype:pdf Example", "search_engine": "test", "rank": 9, "title": "Example research PDF", "snippet": "Example analysis", "normalized_url": "https://papers.test/example.pdf", "source_type": "pdf_document"}),
            ]
            repository.save_collector_result("OSINT_DOCS", CollectorResult("search", "COMPLETED", observations))
            self.assertEqual(repository.get_sources("OSINT_DOCS")[0]["source_url"], "https://papers.test/example.pdf")
            document_sources = repository.get_document_sources("OSINT_DOCS")
            self.assertEqual(document_sources[0]["source_url"], "https://papers.test/example.pdf")
            self.assertEqual(document_sources[0]["discovery_queries"][0]["query_id"], "B001")
            capture_tasks = repository.get_document_capture_tasks("OSINT_DOCS")
            self.assertEqual([item["source_url"] for item in capture_tasks], ["https://papers.test/example.pdf"])

    def test_stores_document_content_and_query_provenance(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = OSINTRepository(Path(temp_dir) / "osint.db")
            target = DomainNormalizer.normalize("example.com")
            repository.create_investigation("OSINT_DOCUMENT_BYTES", target)
            artifact = Path(temp_dir) / "example.pdf"
            artifact.write_bytes(b"verified-pdf-content")
            repository.save_collector_result(
                "OSINT_DOCUMENT_BYTES",
                CollectorResult(
                    "public_search_results",
                    "COMPLETED",
                    [
                        Observation(
                            "public_search_results",
                            "Documents",
                            "PUBLIC_DOCUMENT",
                            "https://files.test/example.pdf",
                            "https://search.test/result",
                            metadata={
                                "artifact_path": str(artifact),
                                "sha256": "sha-test",
                                "content_length": artifact.stat().st_size,
                                "file_name": "example.pdf",
                                "media_type": "application/pdf",
                                "discovery_queries": [{"query_id": "B001", "query_text": "filetype:pdf Example"}],
                            },
                        )
                    ],
                ),
            )
            document = repository.get_documents("OSINT_DOCUMENT_BYTES")[0]
            self.assertEqual(document["file_name"], "example.pdf")
            self.assertEqual(document["discovery_queries"][0]["query_id"], "B001")
            self.assertEqual(
                repository.get_document_content("OSINT_DOCUMENT_BYTES", document["id"]),
                b"verified-pdf-content",
            )

    def test_persists_screenshot_metadata_and_query_metrics(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = OSINTRepository(Path(temp_dir) / "osint.db")
            target = DomainNormalizer.normalize("example.com")
            repository.create_investigation("OSINT_EVIDENCE", target)
            repository.save_queries("OSINT_EVIDENCE", DorkGenerator().generate(target))
            repository.save_collector_result(
                "OSINT_EVIDENCE",
                CollectorResult(
                    "brave_search", "COMPLETED", [Observation(
                        "brave_search", "Search", "SEARCH_RESULT", "Complaint",
                        "https://reviews.example.net/item?utm_source=x", metadata={
                            "query_id": "Q029", "query_text": "query", "search_engine": "brave",
                            "rank": 1, "normalized_url": "https://reviews.example.net/item",
                            "source_type": "review_site", "title": "Complaint",
                        },
                    )]
                ),
            )
            task = repository.get_evidence_tasks("OSINT_EVIDENCE", 5)[0]
            screenshot = Path(temp_dir) / "evidence.png"
            screenshot.write_bytes(b"png")
            record = PageCaptureRecord(
                source_id=task["source_id"], source_url=task["source_url"], final_url=task["source_url"],
                page_title="Complaint", http_status=200, accessibility_status="evidence_found",
                screenshots=[EvidenceScreenshotRecord(
                    "Q029", "External evidence", "External", "brave", 1,
                    ["deposit", "deposit not credited"], ["deposit not credited"],
                    "Deposit not credited", "Deposit was not credited to my wallet.",
                    "exact_phrase", str(screenshot), "abc123", 0.95,
                )],
            )
            repository.save_page_captures("OSINT_EVIDENCE", [record])
            self.assertEqual(repository.get_evidence("OSINT_EVIDENCE")[0]["query_id"], "Q029")
            metrics = {item["query_id"]: item for item in repository.get_query_metrics("OSINT_EVIDENCE")}
            self.assertEqual(metrics["Q029"]["screenshots"], 1)

    def test_exposes_public_document_viewer_captures_in_documents(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = OSINTRepository(Path(temp_dir) / "osint.db")
            target = DomainNormalizer.normalize("example.com")
            repository.create_investigation("OSINT_VIEWER", target)
            repository.save_queries("OSINT_VIEWER", DorkGenerator().generate(target))
            repository.save_collector_result(
                "OSINT_VIEWER",
                CollectorResult(
                    "duckduckgo", "COMPLETED", [Observation(
                        "duckduckgo", "Search", "SEARCH_RESULT", "Example PDF viewer",
                        "https://viewer.example.net/example", metadata={
                            "query_id": "B001", "query_text": 'filetype:pdf "Example"',
                            "search_engine": "duckduckgo", "rank": 1,
                            "normalized_url": "https://viewer.example.net/example",
                            "source_type": "web_page", "title": "Example PDF viewer",
                        },
                    )],
                ),
            )
            task = repository.get_evidence_tasks("OSINT_VIEWER", 5)[0]
            self.assertEqual(task["document_priority"], 1)
            self.assertEqual(task["queries"][0]["document_type"], "pdf")
            screenshot = Path(temp_dir) / "viewer-page.png"
            screenshot.write_bytes(b"png")
            record = PageCaptureRecord(
                source_id=task["source_id"], source_url=task["source_url"],
                final_url=task["source_url"], page_title="Example viewer", http_status=200,
                accessibility_status="document_viewer_captured",
                screenshots=[EvidenceScreenshotRecord(
                    "B001", "Main-brand public PDFs", "documents", "duckduckgo", 1,
                    ["Example"], [], "Public document viewer page/segment 1",
                    "Accessible HTML viewer.", "document_viewer_page", str(screenshot),
                    "viewer-sha", 0.70, "Example", document_page_number=1,
                )],
            )
            repository.save_page_captures("OSINT_VIEWER", [record])
            captures = repository.get_document_viewer_captures("OSINT_VIEWER")
            self.assertEqual(len(captures), 1)
            self.assertEqual(captures[0]["query_id"], "B001")
            self.assertEqual(captures[0]["screenshot_path"], str(screenshot))


class EvidenceCaptureTests(unittest.TestCase):
    def test_every_query_includes_gambling_evidence_vocabulary(self):
        queries = DorkGenerator().generate(DomainNormalizer.normalize("example.com"))
        for query in queries:
            keywords = {item.casefold() for item in query.evidence_keywords}
            self.assertTrue({"betting", "gambling", "casino", "sportsbook", "wagering"}.issubset(keywords))

    def test_groups_nearby_matches_and_sanitizes_paths(self):
        groups = group_evidence_positions([{"documentY": 20}, {"documentY": 200}, {"documentY": 900}])
        self.assertEqual([len(group) for group in groups], [2, 1])
        self.assertEqual(sanitize_path_component("../../bad source:name"), "bad_source_name")

    def test_classifies_manual_review_without_bypass(self):
        self.assertEqual(classify_page_access(403, "")[0], "manual_required")
        self.assertEqual(classify_page_access(200, "Please verify you are human")[0], "manual_required")
        self.assertEqual(classify_page_access(404, "")[0], "failed")

    @patch("osint.documents.extract_pdf_pages")
    def test_pdf_requires_target_and_classifies_relevant_pages(self, extract_pages):
        extract_pages.return_value = [
            "Parimatch payment guide. A deposit can be made using UPI.",
            "Parimatch withdrawal and wallet information.",
        ]
        target = DomainNormalizer.normalize("parimatch.com")
        target = type(target)(target.raw_input, target.domain, target.url, brand="Parimatch")
        result = assess_pdf(b"pdf", target, ["deposit", "withdrawal", "wallet"])
        self.assertTrue(result.accepted)
        self.assertEqual(result.document_type, "Payment / Deposit / Withdrawal")
        self.assertEqual(result.relevant_pages, (1, 2))
        self.assertEqual(result.page_count, 2)

    @patch("osint.documents.extract_pdf_pages")
    def test_pdf_rejects_generic_deposit_document(self, extract_pages):
        extract_pages.return_value = ["A general definition of a bank deposit and payment."]
        target = DomainNormalizer.normalize("parimatch.com")
        result = assess_pdf(b"pdf", target, ["deposit"])
        self.assertFalse(result.accepted)

    @patch("osint.documents.extract_pdf_pages")
    def test_pdf_with_target_only_is_kept_as_a_document(self, extract_pages):
        extract_pages.return_value = ["Independent research paper referencing Parimatch in a market overview."]
        target = DomainNormalizer.normalize("parimatch.com")
        result = assess_pdf(b"pdf", target, ["withdrawal"])
        self.assertTrue(result.accepted)
        self.assertEqual(result.matched_keywords, ())
        self.assertEqual(result.relevant_pages, (1,))

    def test_docx_requires_target_and_extracts_evidence(self):
        buffer = BytesIO()
        document = Document()
        document.add_heading("Diuwin payment guide", level=1)
        document.add_paragraph("Diuwin customers can use a wallet for deposits and withdrawals.")
        document.save(buffer)
        target = DomainNormalizer.normalize("diuwin.example")
        target = type(target)(target.raw_input, target.domain, target.url, brand="Diuwin")
        result = assess_docx(buffer.getvalue(), target, ["wallet", "deposit", "withdrawal"])
        self.assertTrue(result.accepted)
        self.assertEqual(result.document_type, "Payment / Deposit / Withdrawal")

    def test_playwright_highlights_fixture_and_takes_screenshot(self):
        async def exercise():
            resources = await launch_browser_session(headless=True)
            try:
                page = await resources.context.new_page()
                await page.set_content(
                    "<html><body><h1>Fun88 complaint</h1><p>Deposit of Rs 2500 was not credited to the player's wallet.</p></body></html>"
                )
                matches = await SERPEvidenceCapturePipeline.highlight_page(
                    page, ["deposit", "not credited", "wallet"]
                )
                with tempfile.TemporaryDirectory() as temp_dir:
                    path = Path(temp_dir) / "fixture.png"
                    await page.screenshot(path=str(path))
                    self.assertGreater(path.stat().st_size, 0)
                return matches
            finally:
                await close_browser_session(resources)

        try:
            matches = asyncio.run(exercise())
        except Exception as exc:
            self.skipTest(f"Playwright browser unavailable: {exc}")
        self.assertEqual(
            {item["keyword"].casefold() for item in matches},
            {"deposit", "not credited", "wallet"},
        )


class KeylessSearchProviderTests(unittest.TestCase):
    def test_search_collection_stops_at_cancellation_checkpoint(self):
        query = DorkQuery("B001", "documents", "PDFs", "critical", 'filetype:pdf "Example"', "PDFs")
        context = CollectorContext(
            [query],
            search_query_budget=1,
            cancel_check=lambda: True,
        )
        with self.assertRaises(InvestigationCancelled):
            KeylessSearchCollector().collect(DomainNormalizer.normalize("example.com"), context)

    def test_combined_collector_records_each_provider_and_partial_failure(self):
        class StaticProvider(SearchProvider):
            capabilities = BingRSSSearchProvider.capabilities

            def __init__(self, name, *, fail=False):
                self.name = name
                self.fail = fail

            @property
            def available(self):
                return True

            def search(self, query, *, query_id, count):
                if self.fail:
                    raise RuntimeError("temporarily unavailable")
                return [SearchResult(
                    query_id,
                    query,
                    self.name,
                    1,
                    "Example official login",
                    "https://example.com/login",
                    "Example account login",
                )]

        provider = AggregatingSearchProvider([
            StaticProvider("bing_rss"),
            StaticProvider("duckduckgo", fail=True),
        ], name="keyless_aggregated")
        collector = KeylessSearchCollector(provider)
        query = DorkQuery("Q001", "authentication", "Login", "high", '"example" login', "Login pages")

        observations = collector.collect(
            DomainNormalizer.normalize("example.com"),
            CollectorContext([query], search_query_budget=1, results_per_query=5),
        )

        executions = [item for item in observations if item.entity_type == "QUERY_EXECUTION"]
        self.assertEqual(
            {(item.metadata["provider"], item.metadata["status"]) for item in executions},
            {("bing_rss", "completed"), ("duckduckgo", "failed")},
        )
        self.assertTrue(any(item.entity_type == "SEARCH_PROVIDER_ERROR" for item in observations))
        self.assertTrue(any(item.entity_type == "SEARCH_RESULT" for item in observations))

    @patch("osint.search.requests.get")
    def test_parses_google_brand_results(self, get):
        response = Mock()
        response.url = "https://www.google.com/search?q=parimatch"
        response.text = """
        <html><body><div><a href="https://www.scribd.com/document/123"><h3>Parimatch Deposit Guide PDF</h3></a>
        <div class="VwiC3b">Parimatch deposit and withdrawal information.</div></div></body></html>
        """
        response.raise_for_status.return_value = None
        get.return_value = response
        provider = GoogleSearchProvider()
        provider.cache_ttl = 0
        results = provider.search('filetype:pdf "Parimatch"', query_id="B001", count=10)
        self.assertEqual(results[0].url, "https://www.scribd.com/document/123")
        self.assertEqual(results[0].search_engine, "google")

    @patch("osint.search.requests.get")
    def test_google_empty_page_is_not_reported_as_success(self, get):
        response = Mock()
        response.url = "https://www.google.com/search?q=parimatch"
        response.text = "<html><head><title>Before you continue</title></head><body>Consent</body></html>"
        response.raise_for_status.return_value = None
        get.return_value = response
        provider = GoogleSearchProvider()
        provider.api_key = ""
        provider.cse_id = ""
        provider.cache_ttl = 0
        with self.assertRaisesRegex(RuntimeError, "no parseable organic results"):
            provider.search('filetype:pdf "Parimatch"', query_id="B001", count=10)

    @patch("osint.search.requests.get")
    def test_google_429_is_reported_without_http_traceback(self, get):
        response = Mock()
        response.status_code = 429
        response.url = "https://www.google.com/sorry/index"
        response.text = "Too Many Requests"
        get.return_value = response
        provider = GoogleSearchProvider()
        provider.api_key = ""
        provider.cse_id = ""
        provider.cache_ttl = 0
        with self.assertRaisesRegex(RuntimeError, "rate-limited"):
            provider.search('"Tashanwin"', query_id="RESOLVE_01", count=10)
        response.raise_for_status.assert_not_called()
    @patch("osint.search.requests.get")
    def test_parses_duckduckgo_html_results_and_redirect_url(self, get):
        response = Mock()
        response.text = """
        <div class="result">
          <a class="result__a" href="//duckduckgo.com/l/?uddg=https%3A%2F%2Freviews.example%2Fcomplaint">Complaint report</a>
          <a class="result__snippet">Deposit was not credited.</a>
        </div>
        """
        response.raise_for_status.return_value = None
        get.return_value = response
        provider = DuckDuckGoSearchProvider()
        provider.cache_ttl = 0
        results = provider.search("example complaint", query_id="Q029", count=10)
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0].url, "https://reviews.example/complaint")
        self.assertEqual(results[0].search_engine, "duckduckgo")

    @patch("osint.search.requests.get")
    def test_parses_bing_rss_results(self, get):
        response = Mock()
        response.content = b"""<?xml version='1.0'?><rss><channel><item>
          <title>Complaint report</title><link>https://reviews.example/complaint</link>
          <description>Deposit was not credited.</description>
        </item></channel></rss>"""
        response.raise_for_status.return_value = None
        get.return_value = response
        provider = BingRSSSearchProvider()
        provider.cache_ttl = 0
        results = provider.search("example complaint", query_id="Q029", count=10)
        self.assertEqual(results[0].url, "https://reviews.example/complaint")
        self.assertEqual(results[0].search_engine, "bing_rss")

    def test_social_queries_produce_manual_review_links(self):
        target = DomainNormalizer.normalize("example.com")
        query = DorkGenerator().generate(target)[0]
        query = next(item for item in DorkGenerator().generate(target) if item.category == "support_social")
        links = KeylessSearchCollector._manual_social_links(target, query)
        self.assertEqual({item.value for item in links}, {"X/Twitter", "Reddit", "Instagram", "Facebook", "Telegram"})
        self.assertTrue(all(item.metadata["status"] == "manual_required" for item in links))

    def test_social_search_result_becomes_automated_finding(self):
        class SocialProvider(SearchProvider):
            name = "test_social"
            capabilities = BingRSSSearchProvider.capabilities

            @property
            def available(self):
                return True

            def search(self, query, *, query_id, count):
                return [SearchResult(
                    query_id, query, self.name, 2, "Example complaint discussion",
                    "https://www.reddit.com/r/reviews/comments/example",
                    "Users discuss example.com withdrawal complaints.",
                )]

        target = DomainNormalizer.normalize("example.com")
        query = next(item for item in DorkGenerator().generate(target) if item.query_id == "Q022")
        observations = KeylessSearchCollector(SocialProvider()).collect(
            target, CollectorContext([query], search_query_budget=1, results_per_query=5)
        )
        finding = next(item for item in observations if item.entity_type == "AUTOMATED_SOCIAL_FINDING")
        self.assertEqual(finding.metadata["platform"], "Reddit")
        self.assertEqual(finding.metadata["collector_method"], "keyless_public_search")

        with tempfile.TemporaryDirectory() as temp_dir:
            repository = OSINTRepository(Path(temp_dir) / "osint.db")
            repository.create_investigation("OSINT_SOCIAL", target)
            repository.save_collector_result(
                "OSINT_SOCIAL", CollectorResult("keyless_web_search", "COMPLETED", observations)
            )
            stored = repository.get_social_findings("OSINT_SOCIAL")
            self.assertEqual(len(stored), 1)
            self.assertEqual(stored[0]["platform"], "Reddit")
            self.assertEqual(stored[0]["search_rank"], 2)


class FreeDiscoveryTests(unittest.TestCase):
    @patch("osint.collectors.free_discovery.requests.get")
    def test_certificate_transparency_returns_leads_only(self, get):
        response = Mock()
        response.json.return_value = [{"name_value": "www.examplebet.com\nmirror.examplebet.com"}]
        response.raise_for_status.return_value = None
        get.return_value = response
        result = CertificateTransparencyCollector().collect(DomainNormalizer.normalize("examplebet.com"), CollectorContext([]))
        self.assertTrue(result)
        self.assertTrue(all(item.entity_type == "CANDIDATE_DOMAIN" and item.metadata["lead_only"] for item in result))

    @patch("osint.collectors.free_discovery.requests.get")
    def test_wayback_urls_must_pass_the_serp_gate(self, get):
        response = Mock()
        response.json.return_value = [["timestamp", "original"], ["20250101000000", "https://examplebet.com/payments"]]
        response.raise_for_status.return_value = None
        get.return_value = response
        result = WaybackCDXCollector().collect(DomainNormalizer.normalize("examplebet.com"), CollectorContext([], results_per_query=5))
        self.assertEqual(result[0].entity_type, "SEARCH_RESULT")
        self.assertEqual(result[0].metadata["search_engine"], "wayback_cdx")

    def test_docx_export_is_a_valid_office_document(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = OSINTRepository(Path(temp_dir) / "osint.db")
            repository.create_investigation("OSINT_DOCX", DomainNormalizer.normalize("example.com"))
            self.assertTrue(OSINTDocxReportBuilder(repository).build("OSINT_DOCX").startswith(b"PK"))

    def test_docx_includes_each_authenticated_social_screenshot_once(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            repository = OSINTRepository(Path(temp_dir) / "osint.db")
            target = DomainNormalizer.normalize("example.com")
            repository.create_investigation("OSINT_SOCIAL_DOCX", target)
            screenshot = Path(temp_dir) / "x-scroll.png"
            screenshot.write_bytes(base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
            ))
            observations = []
            for rank in (1, 2):
                observations.append(Observation(
                    "x_authenticated_playwright", "Automated social and review findings",
                    "AUTOMATED_SOCIAL_FINDING", f"Example post {rank}",
                    f"https://x.com/example/status/{rank}", confidence=0.9,
                    metadata={
                        "platform": "X/Twitter", "title": f"Example post {rank}",
                        "post_text": f"Target-matching example post {rank}",
                        "matched_target_variant": "example", "query_id": "X-DIRECT",
                        "search_engine": "x_authenticated_browser", "rank": rank,
                        "status": "authenticated_page_captured",
                        "screenshot_paths": [str(screenshot)],
                    },
                ))
            repository.save_collector_result(
                "OSINT_SOCIAL_DOCX",
                CollectorResult("x_authenticated_playwright", "COMPLETED", observations),
            )
            report = Document(BytesIO(OSINTDocxReportBuilder(repository).build("OSINT_SOCIAL_DOCX")))
            self.assertEqual(len(report.inline_shapes), 1)
            self.assertIn(
                "Social and review captures",
                "\n".join(paragraph.text for paragraph in report.paragraphs),
            )

    def test_report_does_not_call_external_brand_evidence_exact_domain(self):
        self.assertEqual(
            evidence_scope(
                "https://pari-match.pro.in/article",
                "https://pari-match.pro.in/article",
                "pari-pro.in",
            ),
            "External brand-related public page",
        )
        self.assertEqual(
            evidence_scope("https://pari-pro.in/", "https://parimatchs123.com/", "pari-pro.in"),
            "Target-domain URL redirected to an external page",
        )

    def test_evidence_text_is_deduplicated_and_not_cut_mid_word(self):
        repeated = (
            "Parimatch supports deposits. Withdrawals are processed quickly.\n\n"
            "Parimatch supports deposits. Withdrawals are processed quickly."
        )
        self.assertEqual(
            clean_evidence_text(repeated),
            "Parimatch supports deposits. Withdrawals are processed quickly.",
        )
        shortened = clean_evidence_text("A complete sentence. " + "word " * 50, limit=45)
        self.assertFalse(shortened.endswith(" wor…"))


if __name__ == "__main__":
    unittest.main()
