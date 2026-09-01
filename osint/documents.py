from __future__ import annotations

import shutil
import subprocess
import tempfile
import hashlib
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from osint.models import NormalizedTarget
from osint.relevance import build_target_variants, find_target_reference, target_keyword_proximity
from osint.text_cleanup import clean_evidence_text


@dataclass(frozen=True)
class DocumentAssessment:
    accepted: bool
    document_type: str = "Other Target-Related Document"
    matched_target_variant: str = ""
    matched_keywords: tuple[str, ...] = ()
    relevant_pages: tuple[int, ...] = ()
    evidence_context: str = ""
    reason: str = ""
    page_count: int = 0


DOCUMENT_CATEGORIES = (
    ("Registration / Company", ("company registration", "certificate of incorporation", "registered office", "company number")),
    ("Licence / Certificate", ("licence", "license", "licensed by", "certificate", "certification")),
    ("Legal / Regulatory", ("court order", "legal notice", "blocking order", "regulator", "regulatory")),
    ("Terms / Policies", ("terms and conditions", "privacy policy", "withdrawal policy", "deposit policy")),
    ("Payment / Deposit / Withdrawal", ("deposit", "withdrawal", "wallet", "payment", "upi", "payout")),
    ("Application / APK", ("apk", "android app", "mobile application", "ios", "ipa")),
    ("Bonus / Promotion", ("welcome bonus", "deposit bonus", "cashback", "free bet", "promotion")),
    ("Review / Analysis", ("review", "complaint", "analysis", "scam", "fraud")),
)


def extract_pdf_pages(content: bytes, timeout: int = 20) -> list[str]:
    """Extract page-delimited text with the local Poppler utility and bounded runtime."""
    executable = shutil.which("pdftotext")
    if not executable:
        return []
    with tempfile.TemporaryDirectory(prefix="osint_pdf_") as temp_dir:
        source = Path(temp_dir) / "source.pdf"
        source.write_bytes(content)
        completed = subprocess.run(
            [executable, "-layout", str(source), "-"],
            capture_output=True,
            check=False,
            timeout=timeout,
        )
    if completed.returncode != 0:
        return []
    return completed.stdout.decode("utf-8", errors="replace").split("\f")


def assess_pdf(content: bytes, target: NormalizedTarget, evidence_keywords: list[str] | tuple[str, ...]) -> DocumentAssessment:
    pages = extract_pdf_pages(content)
    if not pages:
        return DocumentAssessment(False, reason="PDF text could not be extracted; manual review required")
    return assess_document_pages(pages, target, evidence_keywords)


def extract_docx_text(content: bytes) -> list[str]:
    """Extract paragraphs and table cells from a bounded DOCX payload."""
    try:
        from docx import Document

        document = Document(BytesIO(content))
    except Exception:
        return []
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.extend(cell.text for cell in row.cells if cell.text.strip())
    return ["\n".join(blocks)] if blocks else []


def assess_docx(
    content: bytes,
    target: NormalizedTarget,
    evidence_keywords: list[str] | tuple[str, ...],
) -> DocumentAssessment:
    pages = extract_docx_text(content)
    if not pages:
        return DocumentAssessment(False, reason="DOCX text could not be extracted; manual review required")
    return assess_document_pages(pages, target, evidence_keywords)


def assess_document_pages(
    pages: list[str],
    target: NormalizedTarget,
    evidence_keywords: list[str] | tuple[str, ...],
) -> DocumentAssessment:
    """Apply the same target and evidence checks to any extracted document text."""
    variants = build_target_variants(target)
    target_hits = [(number, find_target_reference(text, variants)) for number, text in enumerate(pages, 1)]
    target_hits = [(number, variant) for number, variant in target_hits if variant]
    if not target_hits:
        return DocumentAssessment(False, reason="No verified target reference found in extracted PDF text")

    keywords = tuple(dict.fromkeys(keyword for keyword in evidence_keywords if keyword))
    matches: list[tuple[int, str, str, int]] = []
    contexts = []
    for number, text in enumerate(pages, 1):
        normalized = " ".join(text.split())
        proximity = target_keyword_proximity(normalized, variants, keywords)
        if not proximity:
            continue
        target_variant, keyword, distance = proximity
        lowered = normalized.casefold()
        keyword_index = lowered.find(keyword.casefold())
        start = max(keyword_index - 180, 0)
        end = min(keyword_index + len(keyword) + 300, len(normalized))
        contexts.append(f"Page {number}: {normalized[start:end]}")
        matches.append((number, target_variant, keyword, distance))

    full_text = "\n".join(pages).casefold()
    document_type = next(
        (category for category, terms in DOCUMENT_CATEGORIES if any(term in full_text for term in terms)),
        "Other Target-Related Document",
    )
    if keywords and not matches:
        first_page, first_variant = target_hits[0]
        text = " ".join(pages[first_page - 1].split())
        index = text.casefold().find(first_variant.casefold())
        context = text[max(0, index - 180):index + len(first_variant) + 360]
        return DocumentAssessment(
            True,
            document_type=document_type,
            matched_target_variant=first_variant,
            relevant_pages=(first_page,),
            evidence_context=clean_evidence_text(f"Page {first_page}: {context}", limit=12_000),
            reason="Verified target reference; no configured evidence term occurred within the proximity limit",
            page_count=len(pages),
        )
    return DocumentAssessment(
        True,
        document_type=document_type,
        matched_target_variant=(matches[0][1] if matches else target_hits[0][1]),
        matched_keywords=tuple(sorted({item[2] for item in matches}, key=str.casefold)),
        relevant_pages=tuple(sorted({item[0] for item in matches} or {target_hits[0][0]})),
        evidence_context=clean_evidence_text("\n\n".join(contexts), limit=12_000),
        reason="Verified target reference and target-specific document evidence",
        page_count=len(pages),
    )


def render_pdf_pages(pdf_path: str | Path, page_numbers: tuple[int, ...], output_dir: str | Path) -> list[dict]:
    """Render every requested PDF page in one Poppler invocation."""
    executable = shutil.which("pdftoppm")
    if not executable:
        return []
    destination = Path(output_dir)
    destination.mkdir(parents=True, exist_ok=True)
    requested = tuple(sorted({number for number in page_numbers if number > 0}))
    if not requested:
        return []
    prefix = destination / "page"
    completed = subprocess.run(
        [
            executable, "-f", str(requested[0]), "-l", str(requested[-1]),
            "-png", "-r", "144", str(pdf_path), str(prefix),
        ],
        capture_output=True,
        check=False,
        timeout=max(30, len(requested) * 10),
    )
    captures = []
    if completed.returncode != 0:
        return captures
    requested_set = set(requested)
    for image_path in sorted(destination.glob("page-*.png")):
        try:
            page_number = int(image_path.stem.rsplit("-", 1)[1])
        except (IndexError, ValueError):
            continue
        if page_number not in requested_set:
            continue
        content = image_path.read_bytes()
        captures.append({
            "page": page_number,
            "path": str(image_path),
            "sha256": hashlib.sha256(content).hexdigest(),
        })
    return captures
