from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches

from osint.storage import OSINTRepository
from osint.text_cleanup import clean_evidence_text, evidence_scope


class OSINTDocxReportBuilder:
    """Build a portable report containing only stored, target-validated evidence."""

    def __init__(self, repository: OSINTRepository):
        self.repository = repository

    def build(self, investigation_id: str) -> bytes:
        investigation = self.repository.get_investigation(investigation_id)
        evidence = self.repository.get_evidence(investigation_id)
        documents = self.repository.get_documents(investigation_id)
        domains = self.repository.get_candidates(investigation_id)
        report = Document()
        report.add_heading("Web intelligence and OSINT report", 0)
        report.add_paragraph(f"Investigation: {investigation_id}")
        report.add_heading("Target", level=1)
        for label, value in (
            ("Original input", investigation.get("original_input")),
            ("Resolved domain", investigation.get("target_domain")),
            ("Resolved brand", investigation.get("resolved_brand")),
            ("Status", investigation.get("status")),
        ):
            report.add_paragraph(f"{label}: {value or 'Unknown'}")

        report.add_heading("Resolved / related domains", level=1)
        report.add_paragraph("Traffic figures are third-party estimated visits, not exact site analytics.")
        table = report.add_table(rows=1, cols=7)
        for cell, label in zip(table.rows[0].cells, ("Domain", "Status", "HTTP", "Final URL", "Monthly Visits", "Yearly Visits", "Traffic source")):
            cell.text = label
        for item in domains:
            cells = table.add_row().cells
            values = (
                item.get("domain"), item.get("domain_status", "Unknown"), item.get("http_status", "Unavailable"),
                item.get("final_url") or "", item.get("monthly_visits") if item.get("monthly_visits") is not None else "Unavailable",
                (str(item.get("yearly_visits")) + (" (Projected)" if item.get("yearly_visits_kind") == "Projected" else "")) if item.get("yearly_visits") is not None else "Unavailable",
                item.get("traffic_source", "Unavailable"),
            )
            for cell, value in zip(cells, values):
                cell.text = str(value)

        report.add_heading("Verified web evidence", level=1)
        if not evidence and not documents:
            report.add_paragraph("No confirmed target-specific public evidence found.")
        for item in evidence:
            scope = evidence_scope(
                item["source_url"], item.get("final_url") or item["source_url"],
                investigation.get("target_domain") or "",
            )
            report.add_heading(f"Web evidence — {scope}", level=2)
            report.add_paragraph(f"Discovery query: {item.get('query_name') or item['query_id']}")
            report.add_paragraph(f"Source URL: {item['source_url']}")
            report.add_paragraph(f"Final URL: {item.get('final_url') or item['source_url']}")
            report.add_paragraph(f"Matched target: {item.get('matched_target_variant') or 'stored page match'}")
            report.add_paragraph(f"Matched keywords: {', '.join(item.get('matched_keywords') or [])}")
            report.add_paragraph(clean_evidence_text(item.get("context_text") or item.get("evidence_text")))
            image_path = Path(item["screenshot_path"])
            if image_path.is_file():
                report.add_picture(str(image_path), width=Inches(6.2))
                report.add_paragraph(f"Screenshot SHA-256: {item.get('sha256') or ''}")

        if documents:
            report.add_heading("Verified public documents", level=1)
        for item in documents:
            report.add_heading(item.get("document_type") or "Public document", level=2)
            report.add_paragraph(f"Source URL: {item['source_url']}")
            report.add_paragraph(f"Matched target: {item.get('matched_target_variant') or ''}")
            report.add_paragraph(f"Relevant pages: {', '.join(map(str, item.get('relevant_pages') or []))}")
            report.add_paragraph(clean_evidence_text(item.get("evidence_context"), limit=12_000))
            for capture in item.get("page_screenshots", []):
                image_path = Path(capture.get("path", ""))
                if image_path.is_file():
                    report.add_picture(str(image_path), width=Inches(6.2))
                    report.add_paragraph(f"Document page {capture.get('page')} SHA-256: {capture.get('sha256')}")
        output = BytesIO()
        report.save(output)
        return output.getvalue()
